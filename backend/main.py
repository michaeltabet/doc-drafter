import io
import os
import re
import json
import logging
import unicodedata
from typing import Optional

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse

import pdfplumber
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image
from docx import Document

# ─── Config ───
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
MAX_OCR_PAGES = 20
OCR_TIMEOUT = 30  # seconds per page
MAX_IMAGE_PIXELS = 25_000_000

Image.MAX_IMAGE_PIXELS = MAX_IMAGE_PIXELS

ALLOWED_ORIGINS = os.environ.get(
    "ALLOWED_ORIGINS",
    "https://drafter.michaeltabet.com"
).split(",")

ALLOWED_EXTENSIONS = {
    "pdf", "docx", "png", "jpg", "jpeg", "tiff", "bmp", "gif", "webp",
    "txt", "md", "html", "htm", "rtf",
}

TEXT_EXTENSIONS = {"txt", "md", "html", "htm", "rtf", "csv", "xml", "json"}

# ─── Logging ───
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(message)s",
)
logger = logging.getLogger("doc-drafter")

# ─── App ───
app = FastAPI(title="Doc Drafter API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_methods=["GET", "POST"],
    allow_headers=["Content-Type"],
)


# ─── Helpers ───

async def read_upload(file: UploadFile, max_size: int = MAX_FILE_SIZE) -> bytes:
    """Read an uploaded file in chunks, enforcing size limit before full read."""
    chunks = []
    total = 0
    while True:
        chunk = await file.read(64 * 1024)
        if not chunk:
            break
        total += len(chunk)
        if total > max_size:
            raise HTTPException(413, f"File too large (max {max_size // (1024*1024)}MB)")
        chunks.append(chunk)
    return b"".join(chunks)


def validate_extension(filename: str) -> str:
    """Validate file extension and return it."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            415,
            f"Unsupported file type: .{ext}. Supported: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )
    return ext


def sanitize_filename(name: str) -> str:
    """Sanitize a filename for safe use in Content-Disposition headers."""
    name = unicodedata.normalize("NFKD", name)
    name = re.sub(r'[^\w\s\-.]', '', name)
    name = name.strip()[:200]
    return name or "document"


# ─── Text Extraction ───

def extract_pdf_text(data: bytes) -> tuple[str, list[str]]:
    """Extract text from PDF using pdfplumber, fallback to OCR for scanned pages."""
    warnings = []
    text_parts = []

    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages[:MAX_OCR_PAGES]:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
        if len(pdf.pages) > MAX_OCR_PAGES:
            warnings.append(f"PDF has {len(pdf.pages)} pages, only first {MAX_OCR_PAGES} processed")

    # OCR fallback for pages with no text
    empty_pages = [i for i, t in enumerate(text_parts) if not t.strip()]
    if empty_pages:
        try:
            for page_idx in empty_pages:
                images = convert_from_bytes(
                    data, first_page=page_idx + 1, last_page=page_idx + 1, dpi=150
                )
                if images:
                    ocr_text = pytesseract.image_to_string(images[0], timeout=OCR_TIMEOUT)
                    text_parts[page_idx] = ocr_text
                    images[0].close()
        except Exception as e:
            logger.warning(f"OCR fallback failed: {e}")
            warnings.append(f"OCR failed for scanned pages: {e}")

    return "\n\n".join(text_parts), warnings


def extract_docx_text(data: bytes) -> str:
    """Extract text from DOCX preserving paragraph structure."""
    doc = Document(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text for cell in row.cells)
            parts.append(row_text)
    return "\n".join(parts)


def extract_image_text(data: bytes) -> str:
    """OCR an image file directly."""
    img = Image.open(io.BytesIO(data))
    if img.width * img.height > MAX_IMAGE_PIXELS:
        raise HTTPException(413, "Image too large for processing")
    try:
        return pytesseract.image_to_string(img, timeout=OCR_TIMEOUT)
    finally:
        img.close()


def extract_text(data: bytes, filename: str) -> tuple[str, list[str]]:
    ext = validate_extension(filename)
    warnings = []

    if ext == "pdf":
        return extract_pdf_text(data)
    elif ext == "docx":
        return extract_docx_text(data), warnings
    elif ext in ("png", "jpg", "jpeg", "tiff", "bmp", "gif", "webp"):
        return extract_image_text(data), warnings
    elif ext in TEXT_EXTENSIONS:
        try:
            return data.decode("utf-8"), warnings
        except UnicodeDecodeError:
            raise HTTPException(422, "File does not appear to be valid UTF-8 text")
    else:
        raise HTTPException(415, f"Unsupported file type: .{ext}")


# ─── Placeholder Logic ───

def find_placeholders(text: str, open_delim: str, close_delim: str) -> list[str]:
    if not open_delim or not close_delim:
        raise HTTPException(422, "Delimiters cannot be empty")
    open_esc = re.escape(open_delim)
    close_esc = re.escape(close_delim)
    pattern = open_esc + r"\s*([\w][\w\s.\-]*)\s*" + close_esc
    found = list(dict.fromkeys(m.strip() for m in re.findall(pattern, text[:500_000])))
    return found


def placeholder_to_label(name: str) -> str:
    return re.sub(r"[_\-.]", " ", name).title()


def guess_type(name: str) -> dict:
    n = name.lower()
    if "date" in n:
        return {"type": "string", "format": "date"}
    if "email" in n:
        return {"type": "string", "format": "email"}
    if "phone" in n or "tel" in n:
        return {"type": "string", "format": "tel"}
    if re.search(r"amount|price|total|cost|fee|salary|rate", n):
        return {"type": "number"}
    if re.search(r"count|quantity|num_|number_of", n):
        return {"type": "integer"}
    if re.search(r"address|description|notes|terms|clause|body", n):
        return {"type": "string", "multiline": True}
    return {"type": "string"}


def build_schema(placeholders: list[str], title: str) -> dict:
    properties = {}
    required = []
    for ph in placeholders:
        info = guess_type(ph)
        prop = {"type": info.get("type", "string"), "title": placeholder_to_label(ph)}
        if "format" in info:
            prop["format"] = info["format"]
        if info.get("multiline"):
            prop["_multiline"] = True
        properties[ph] = prop
        required.append(ph)

    return {
        "$schema": "http://json-schema.org/draft-07/schema#",
        "type": "object",
        "title": title,
        "properties": properties,
        "required": required,
    }


# ─── DOCX Filling ───

def fill_docx(data: bytes, open_delim: str, close_delim: str, form_data: dict) -> bytes:
    """Fill a DOCX template preserving formatting."""
    doc = Document(io.BytesIO(data))

    def replace_in_paragraph(paragraph):
        full_text = "".join(run.text for run in paragraph.runs)
        has_placeholder = any(
            f"{open_delim}{key}{close_delim}" in full_text for key in form_data
        )
        if not has_placeholder:
            return

        for key, value in form_data.items():
            placeholder = f"{open_delim}{key}{close_delim}"
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, str(value or ""))

        if paragraph.runs:
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ""

    for para in doc.paragraphs:
        replace_in_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)

    for section in doc.sections:
        for header_para in section.header.paragraphs:
            replace_in_paragraph(header_para)
        for footer_para in section.footer.paragraphs:
            replace_in_paragraph(footer_para)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── Endpoints ───

@app.post("/api/parse")
async def parse_template(
    file: UploadFile = File(...),
    open_delim: str = Form("{{"),
    close_delim: str = Form("}}"),
):
    data = await read_upload(file)

    filename = file.filename or "document"
    validate_extension(filename)

    logger.info(f"Parsing: {filename} ({len(data)} bytes)")
    text, warnings = extract_text(data, filename)

    if not text.strip():
        raise HTTPException(422, "Could not extract any text from this file")

    placeholders = find_placeholders(text, open_delim, close_delim)
    title = filename.rsplit(".", 1)[0] if "." in filename else filename
    schema = build_schema(placeholders, title) if placeholders else None

    return {
        "text": text,
        "placeholders": placeholders,
        "schema": schema,
        "filename": filename,
        "warnings": warnings,
    }


@app.post("/api/generate")
async def generate_document(
    file: UploadFile = File(...),
    open_delim: str = Form("{{"),
    close_delim: str = Form("}}"),
    form_data: str = Form("{}"),
):
    data = await read_upload(file)

    try:
        parsed_data = json.loads(form_data)
    except json.JSONDecodeError as e:
        raise HTTPException(422, f"Invalid JSON in form_data: {e}")

    if not isinstance(parsed_data, dict):
        raise HTTPException(422, "form_data must be a JSON object")
    if any(not isinstance(v, (str, int, float, bool, type(None))) for v in parsed_data.values()):
        raise HTTPException(422, "form_data values must be scalar types")

    filename = file.filename or "document"
    ext = validate_extension(filename)

    logger.info(f"Generating: {filename} ({len(data)} bytes, {len(parsed_data)} fields)")

    safe_name = sanitize_filename(filename.rsplit(".", 1)[0])

    if ext == "docx":
        filled = fill_docx(data, open_delim, close_delim, parsed_data)
        output_name = f"{safe_name}_filled.docx"
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        text, _ = extract_text(data, filename)
        for key, value in parsed_data.items():
            pattern = re.escape(open_delim) + r"\s*" + re.escape(key) + r"\s*" + re.escape(close_delim)
            text = re.sub(pattern, str(value or ""), text)
        filled = text.encode("utf-8")
        output_name = f"{safe_name}_filled.txt"
        media_type = "text/plain"

    return StreamingResponse(
        io.BytesIO(filled),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{output_name}"'},
    )


@app.get("/api/health")
async def health():
    return {"status": "ok"}
